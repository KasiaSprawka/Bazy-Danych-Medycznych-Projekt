#!/usr/bin/python
# -*- coding: utf-8 -*-

from docx import Document,styles
from docx.shared import Inches, Pt
#from docx.text.run import Font
from docx.enum.style import WD_STYLE_TYPE, WD_STYLE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import MySQLdb
import time

w_id = 8

# Open database connection
db = MySQLdb.connect("localhost","Kasia :)","mysql","protetyka_dentystyczna")

# prepare a cursor object using cursor() method
cursor = db.cursor()

cursor.execute("SET NAMES utf8")

query = """
    SELECT 
    CONCAT(o1.imie, ' ', o1.nazwisko) AS pacjent,
    p.pesel AS pesel,
    CONCAT(o2.imie, ' ', o2.nazwisko) AS lekarz,
    w.data_wizyty
FROM
    wizyty w
        INNER JOIN
    osoby o1 ON o1.id = w.p_id
        INNER JOIN
    osoby o2 ON o2.id = w.l_id
        JOIN
    pacjenci p ON w.p_id = p.p_id
WHERE
    w.w_id =
""" + str(w_id)+';'

# execute SQL query using execute() method.
cursor.execute(query)

# Fetch a single row using fetchone() method.
dane = cursor.fetchone()
 
db.close()

zgoda = Document('zgoda_template.docx')

styles = zgoda.styles
style = styles.add_style('Pogrubienie', WD_STYLE_TYPE.PARAGRAPH)
style.base_style = styles['Normal']
style = zgoda.styles['Pogrubienie']
font = style.font
font.name = 'Times NewRoman'
font.size = Pt(11.5)
font.bold = True


style = zgoda.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10.5)
font.bold = False


p=zgoda.add_paragraph(unicode('Imię i nazwisko pacjenta, nr PESEL: %s, %s\n' %(dane[0], dane[1]),'utf-8'),style = 'Pogrubienie')
p2 = zgoda.add_paragraph(unicode("""Zgodnie z art. 32 – 35 ustawy z dnia 5 grudnia 1996 r. o zawodach lekarza i lekarza dentysty (tekst jednolity Dz. U. z 2008 nr 136 poz. 857 z późniejszymi zmianami) oraz art. 16 -18 ustawy z dnia 6 listopada 2008 r. o prawach pacjenta i Rzeczniku Praw Pacjenta (Dz. U. 2009 r. Nr 52 poz. 417 z późn. zm.) wyrażam zgodę na leczenie chirurgiczne - wszczepienie implantów zęba/zębów przez lek. dent. %s w Łodzi w dniu %s                         .\n\n""" %(dane[2], dane[3]),'utf-8'), style = 'Normal2')
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p2.line_spacing_rule = WD_LINE_SPACING.SINGLE
p2.add_run(unicode('Oświadczam, że udzieliłem(-am) wyczerpujących i prawdziwych informacji co do mojego stanu zdrowia - zgodnie z ankietą stanowiącą załącznik nr 1 do niniejszego oświadczenia. O wszelkich zmianach stanu mojego zdrowia zobowiązuję się powiadomić lekarza prowadzącego. Przyjmuję do wiadomości, że w/w są danymi poufnymi. Wyrażam zgodę na wykonanie dokumentacji radiologicznej i fotograficznej.', 'utf-8-')).font.size=Pt(10.5)
p3=zgoda.add_paragraph(unicode('Zostałem(am) poinformowany(a):','utf-8'),style = 'Pogrubienie')
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
p3.line_spacing_rule = WD_LINE_SPACING.SINGLE
p4=zgoda.add_paragraph(unicode('o technice zabiegu i szczegółowo zaznajomiony (-a) z całym przebiegiem proponowanego leczenia, a także o tym, że zabieg wykonywany będzie w znieczuleniu miejscowym.','utf-8'),style = 'List Paragraph')
p6=zgoda.add_paragraph(unicode('o tym, że ostateczną decyzję o wszczepieniu implantów lekarz podejmuje dopiero podczas zabiegu chirurgicznego, po odsłonięciu kości wyrostka zębodołowego. Kość może mieć niekorzystną budowę dla wszczepienia implantu, co nie zawsze można stwierdzić na podstawie badania RTG.','utf-8'), style='List Paragraph')
p7=zgoda.add_paragraph(unicode('o ryzyku i możliwości wystąpienia powikłań ( w trakcie lub po zabiegu), takich jak: obrzęk, krwawienie, stany zapalne, utrudnione gojenie się rany operacyjnej i w konsekwencji odsłonięcie wszczepu śródkostnego. O tym, że w przypadku niedostatecznego, domowego utrzymywania higieny jamy ustnej może dojść do stanu zapalnego tkanek wokół implantu, w szczególności do stanu zapalnego kości i w konsekwencji konieczności jego usunięcia.','utf-8'), style='List Paragraph')
p8=zgoda.add_paragraph(unicode('o niekorzystnym wpływie palenia tytoniu na ostateczny wynik leczenia implantologicznego.','utf-8'), style='List Paragraph')
p9=zgoda.add_paragraph(unicode('o wskazaniach odnośnie postępowania po zabiegu, a w szczególności o:','utf-8'), style='List Paragraph')
p11=zgoda.add_paragraph(unicode('zakazie prowadzenia pojazdów mechanicznych przez 12 godzin po zabiegu','utf-8'), style='Heading 3')
p12=zgoda.add_paragraph(unicode('zakazie picia alkoholu i palenia tytoniu przez co najmniej 10 dni po zabiegu','utf-8'), style='Heading 3')
p13=zgoda.add_paragraph(unicode('konieczności przyjmowania przez pacjenta przepisanego antybiotyku','utf-8'), style='Heading 3')
p14=zgoda.add_paragraph(unicode('konieczności płukania jamy ustnej zaleconym antyseptykiem','utf-8'), style='Heading 3')
p15=zgoda.add_paragraph(unicode('konieczności usunięciu szwów chirurgicznych','utf-8'), style='Heading 3')
p16=zgoda.add_paragraph(unicode('wizytach kontrolnych, na które pacjent musi się zgłaszać w wyznaczonych przez lekarza\tterminach','utf-8'), style='Heading 3')
p17=zgoda.add_paragraph(unicode('konieczności ścisłego przestrzegania zaleceń dotyczących higieny jamy ustnej','utf-8'), style='Heading 3')
p10=zgoda.add_paragraph(unicode('o kosztach leczenia, które akceptuję.' ,'utf-8'), style = 'List Paragraph')
p5 = zgoda.add_paragraph(unicode('\nPowyższe zasady przeczytałem/-am i zrozumiałem/-am, uzyskałem/-am również wszelkie wyjaśnienia dotyczące leczenia w moim przypadku. Zostałem/-am poinformowany/-a o alternatywnych możliwościach leczenia, z zaniechaniem leczenia włącznie. Zostałem/-am poinformowany/-a o ryzyku towarzyszącym innym metodom leczenia i konsekwencjach wynikających z zaniechania leczenia. Rozumiem, że tak jak w przypadku wszystkich procedur ogólnomedycznych, pozytywne efekty leczenia nie są zagwarantowane. Ponadto, zabieg wszczepienia implantów jest wykonywany w celu usunięcia konkretnego problemu i może nie wyeliminować innych ukrytych problemów. Leczenie to nie zabezpiecza przed próchnicą oraz chorobami przyzębia. Wiem, że mogę odwołać zgodę na leczenie.                             .','utf-8'))
p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p5.line_spacing_rule = WD_LINE_SPACING.SINGLE
p18=zgoda.add_paragraph(unicode("""\n\n\n\n%s,\n
Data, Podpis i pieczątka lekarza dentysty                         Czytelny podpis pacjenta (Rodzica lub opiekuna)"""%(time.strftime("%d/%m/%Y")) ,'utf-8'), style = 'Normal')
p18.alignment = WD_ALIGN_PARAGRAPH.LEFT

zgoda.save(dane[0]+'.docx')
print 'Gotowe :)'